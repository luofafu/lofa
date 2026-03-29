import docx
from docx import Document

doc = Document('./你好Python.docx')
print("段落数",len(doc.paragraphs))
print(doc.paragraphs[0].text)

# 读出全部内容 会读出空格
for p in doc.paragraphs:
    print(p.text)


for p in doc.paragraphs:
    for run in p.runs:
        print(run.text)

# 读取标题
for p in doc.paragraphs:
    style_name = p.style.name
    # print(style_name)
    if style_name.startswith('Heading'):
        print(style_name)
        print(p.text)
# 修改内容
for i in range(len(doc.paragraphs)):
    for run in doc.paragraphs[i].runs:
        if "Pluto" in run.text:
            run.text = run.text.replace("Pluto","lofa")
doc.save('你好lofa.docx')

# # 创建文档
# doc = Document()
#
# doc.add_heading('你好', 1)
# doc.add_paragraph('我爱Python',style='List Bullet')
# doc.save('新保存文件.docx')
# # 文字格式
# doc.add_run('你好').bold = True